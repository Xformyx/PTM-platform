"""
Markdown to Word (.docx) Converter for PTM Analysis Reports
v45: Academic paper style conversion with proper formatting + base64 image embedding + LaTeX symbols + HTML tag conversion

Converts Markdown report to a professionally formatted Word document
with Times New Roman font, proper heading hierarchy, tables, images, and references.
"""

import re
import os
import base64
import tempfile
import io
from typing import List, Tuple, Optional

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


import logging

_logger = logging.getLogger("ptm-workers.docx-converter")

def sse_log(message, level="INFO"):
    if level == "WARNING":
        _logger.warning(message)
    else:
        _logger.info(message)


def setup_document_styles(doc: 'Document'):
    """Configure academic paper styles for the document"""
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Set paragraph format
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.space_before = Pt(0)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.15
    
    # Title style (for report title)
    title_style = doc.styles['Title']
    title_font = title_style.font
    title_font.name = 'Times New Roman'
    title_font.size = Pt(16)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Heading 1 (## sections like Abstract, Introduction, etc.)
    h1_style = doc.styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Times New Roman'
    h1_font.size = Pt(14)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 0, 0)
    h1_style.paragraph_format.space_before = Pt(18)
    h1_style.paragraph_format.space_after = Pt(8)
    
    # Heading 2 (### subsections)
    h2_style = doc.styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Times New Roman'
    h2_font.size = Pt(12)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0, 0, 0)
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(6)
    
    # Heading 3 (#### sub-subsections)
    h3_style = doc.styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = 'Times New Roman'
    h3_font.size = Pt(11)
    h3_font.bold = True
    h3_font.italic = True
    h3_font.color.rgb = RGBColor(0, 0, 0)
    h3_style.paragraph_format.space_before = Pt(8)
    h3_style.paragraph_format.space_after = Pt(4)
    
    # Set page margins (1 inch = 2.54 cm)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def parse_markdown_line(line: str) -> Tuple[str, List[Tuple[str, dict]]]:
    """
    Parse a markdown line and return (type, segments).
    Segments are (text, format_dict) tuples.
    """
    # Check for headings
    heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
    if heading_match:
        level = len(heading_match.group(1))
        text = heading_match.group(2).strip()
        return (f'heading_{level}', [(text, {})])
    
    # Check for horizontal rule
    if re.match(r'^---+\s*$', line) or re.match(r'^\*\*\*+\s*$', line):
        return ('hr', [])
    
    # Check for image line: ![alt](src)
    img_match = re.match(r'^!\[([^\]]*)\]\((.+)\)\s*$', line.strip())
    if img_match:
        alt_text = img_match.group(1)
        img_src = img_match.group(2)
        return ('image', [(alt_text, {'src': img_src})])
    
    # Check for list items
    list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.*)', line)
    if list_match:
        indent = len(list_match.group(1))
        marker = list_match.group(2)
        text = list_match.group(3)
        is_ordered = bool(re.match(r'\d+\.', marker))
        segments = parse_inline_formatting(text)
        return ('list_item', segments)
    
    # Regular paragraph
    if line.strip():
        segments = parse_inline_formatting(line)
        return ('paragraph', segments)
    
    return ('empty', [])


def _convert_latex_symbols(text: str) -> str:
    """
    v42: Convert common LaTeX notation to Unicode symbols for Word documents.
    Handles $\\beta$, $\\alpha$, $\\gamma$, etc.
    """
    # LaTeX to Unicode mapping
    latex_map = {
        r'$\beta$': '\u03B2', r'$\\beta$': '\u03B2',
        r'$\alpha$': '\u03B1', r'$\\alpha$': '\u03B1',
        r'$\gamma$': '\u03B3', r'$\\gamma$': '\u03B3',
        r'$\delta$': '\u03B4', r'$\\delta$': '\u03B4',
        r'$\epsilon$': '\u03B5', r'$\\epsilon$': '\u03B5',
        r'$\kappa$': '\u03BA', r'$\\kappa$': '\u03BA',
        r'$\lambda$': '\u03BB', r'$\\lambda$': '\u03BB',
        r'$\mu$': '\u03BC', r'$\\mu$': '\u03BC',
        r'$\pi$': '\u03C0', r'$\\pi$': '\u03C0',
        r'$\sigma$': '\u03C3', r'$\\sigma$': '\u03C3',
        r'$\tau$': '\u03C4', r'$\\tau$': '\u03C4',
        r'$\phi$': '\u03C6', r'$\\phi$': '\u03C6',
        r'$\omega$': '\u03C9', r'$\\omega$': '\u03C9',
        r'$\Delta$': '\u0394', r'$\\Delta$': '\u0394',
        r'$\Sigma$': '\u03A3', r'$\\Sigma$': '\u03A3',
    }
    for latex, unicode_char in latex_map.items():
        text = text.replace(latex, unicode_char)
    
    # Generic pattern: $\symbolname$ -> attempt conversion
    import re
    def _latex_replace(m):
        sym = m.group(1).lower()
        greek = {
            'alpha': '\u03B1', 'beta': '\u03B2', 'gamma': '\u03B3', 'delta': '\u03B4',
            'epsilon': '\u03B5', 'zeta': '\u03B6', 'eta': '\u03B7', 'theta': '\u03B8',
            'iota': '\u03B9', 'kappa': '\u03BA', 'lambda': '\u03BB', 'mu': '\u03BC',
            'nu': '\u03BD', 'xi': '\u03BE', 'pi': '\u03C0', 'rho': '\u03C1',
            'sigma': '\u03C3', 'tau': '\u03C4', 'upsilon': '\u03C5', 'phi': '\u03C6',
            'chi': '\u03C7', 'psi': '\u03C8', 'omega': '\u03C9',
        }
        return greek.get(sym, m.group(0))
    text = re.sub(r'\$\\?\\?([a-zA-Z]+)\$', _latex_replace, text)
    
    return text


def _convert_html_tags(text: str) -> str:
    """
    v45: Convert HTML inline tags to Markdown equivalents before processing.
    Handles <i>, <em>, <b>, <strong>, <sub>, <sup>, <br>, and strips unknown tags.
    """
    # Convert italic tags to markdown italic
    text = re.sub(r'<i>(.*?)</i>', r'*\1*', text, flags=re.IGNORECASE)
    text = re.sub(r'<em>(.*?)</em>', r'*\1*', text, flags=re.IGNORECASE)
    # Convert bold tags to markdown bold
    text = re.sub(r'<b>(.*?)</b>', r'**\1**', text, flags=re.IGNORECASE)
    text = re.sub(r'<strong>(.*?)</strong>', r'**\1**', text, flags=re.IGNORECASE)
    # Convert subscript/superscript to plain text (Word doesn't support via markdown)
    text = re.sub(r'<sub>(.*?)</sub>', r'\1', text, flags=re.IGNORECASE)
    text = re.sub(r'<sup>(.*?)</sup>', r'\1', text, flags=re.IGNORECASE)
    # Convert <br> to space
    text = re.sub(r'<br\s*/?>', ' ', text, flags=re.IGNORECASE)
    # Strip any remaining unknown HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    return text


def parse_inline_formatting(text: str) -> List[Tuple[str, dict]]:
    """
    Parse inline markdown formatting (bold, italic, code) and return segments.
    Each segment is (text, format_dict) where format_dict has keys like 'bold', 'italic', 'code'.
    """
    # v45: Convert HTML tags to markdown equivalents first
    text = _convert_html_tags(text)
    # v42: Convert LaTeX symbols
    text = _convert_latex_symbols(text)
    segments = []
    
    # Pattern to match bold+italic, bold, italic, inline code
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'   # bold+italic
        r'|(\*\*(.+?)\*\*)'       # bold
        r'|(\*(.+?)\*)'           # italic
        r'|(`(.+?)`)'             # inline code
    )
    
    last_end = 0
    for match in pattern.finditer(text):
        # Add text before this match
        if match.start() > last_end:
            segments.append((text[last_end:match.start()], {}))
        
        if match.group(2):  # bold+italic
            segments.append((match.group(2), {'bold': True, 'italic': True}))
        elif match.group(4):  # bold
            segments.append((match.group(4), {'bold': True}))
        elif match.group(6):  # italic
            segments.append((match.group(6), {'italic': True}))
        elif match.group(8):  # code
            segments.append((match.group(8), {'code': True}))
        
        last_end = match.end()
    
    # Add remaining text
    if last_end < len(text):
        segments.append((text[last_end:], {}))
    
    if not segments:
        segments = [(text, {})]
    
    return segments


def add_formatted_paragraph(doc: 'Document', segments: List[Tuple[str, dict]], style=None, alignment=None, indent_level=0):
    """Add a paragraph with formatted segments to the document"""
    p = doc.add_paragraph(style=style)
    
    if alignment:
        p.alignment = alignment
    
    if indent_level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * indent_level)
    
    for text, fmt in segments:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        
        if fmt.get('bold'):
            run.bold = True
        if fmt.get('italic'):
            run.italic = True
        if fmt.get('code'):
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            # Light gray background for code
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0"/>')
            run._element.get_or_add_rPr().append(shading)
    
    return p


def parse_markdown_table(lines: List[str], start_idx: int) -> Tuple[List[List[str]], int]:
    """
    Parse a markdown table starting at start_idx.
    Returns (rows, end_idx) where rows is a list of lists of cell strings.
    """
    rows = []
    i = start_idx
    
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        
        # Skip separator row (|---|---|...)
        if re.match(r'^\|[\s\-:]+\|', line):
            i += 1
            continue
        
        # Parse cells
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
        i += 1
    
    return rows, i


def add_table_to_doc(doc: 'Document', rows: List[List[str]]):
    """Add a formatted table to the document"""
    if not rows:
        return
    
    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Style the table
    table.style = 'Table Grid'
    
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = ''
                p = cell.paragraphs[0]
                
                # Parse inline formatting in cell text
                segments = parse_inline_formatting(cell_text)
                for text, fmt in segments:
                    run = p.add_run(text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    
                    if fmt.get('bold'):
                        run.bold = True
                    if fmt.get('italic'):
                        run.italic = True
                
                # Header row styling (first row)
                if i == 0:
                    for run in p.runs:
                        run.bold = True
                    # Gray background for header
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
                    cell._element.get_or_add_tcPr().append(shading)
    
    # Add spacing after table
    doc.add_paragraph()


def decode_base64_image(data_uri: str) -> Optional[bytes]:
    """
    Decode a base64 data URI to raw image bytes.
    Handles formats like: data:image/png;base64,iVBOR...
    Also handles raw base64 strings without the data URI prefix.
    """
    try:
        if data_uri.startswith('data:'):
            # Extract base64 portion after the comma
            _, encoded = data_uri.split(',', 1)
        else:
            encoded = data_uri
        
        return base64.b64decode(encoded)
    except Exception as e:
        sse_log(f"Failed to decode base64 image: {e}", "WARNING")
        return None


def add_image_to_doc(doc: 'Document', img_src: str, alt_text: str = "", max_width_inches: float = 6.0):
    """
    Add an image to the document. Supports:
    - base64 data URIs (data:image/png;base64,...)
    - Local file paths
    
    Args:
        doc: The Document object
        img_src: Image source (base64 data URI or file path)
        alt_text: Alternative text for the image
        max_width_inches: Maximum width in inches (default 6.0 for letter paper with 1" margins)
    
    Returns:
        True if image was added successfully, False otherwise
    """
    try:
        if img_src.startswith('data:'):
            # Base64 encoded image
            image_bytes = decode_base64_image(img_src)
            if not image_bytes:
                sse_log(f"Failed to decode base64 image for: {alt_text}", "WARNING")
                return False
            
            # Write to a temporary file for python-docx
            # Determine extension from data URI
            ext = '.png'  # default
            if 'image/jpeg' in img_src or 'image/jpg' in img_src:
                ext = '.jpg'
            elif 'image/gif' in img_src:
                ext = '.gif'
            elif 'image/svg' in img_src:
                ext = '.svg'
            
            # Use BytesIO stream for python-docx
            image_stream = io.BytesIO(image_bytes)
            
            # Add image with width constraint
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            
            # Try to get image dimensions to maintain aspect ratio
            try:
                from PIL import Image as PILImage
                pil_img = PILImage.open(io.BytesIO(image_bytes))
                img_width, img_height = pil_img.size
                
                # Calculate appropriate width (max 6 inches for standard margins)
                width_inches = min(max_width_inches, img_width / 96.0)  # assume 96 DPI
                if width_inches < 3.0:
                    width_inches = min(max_width_inches, 5.0)  # ensure reasonable size
                
                run.add_picture(image_stream, width=Inches(width_inches))
            except ImportError:
                # PIL not available, use default width
                run.add_picture(image_stream, width=Inches(max_width_inches))
            
            sse_log(f"Embedded base64 image: {alt_text}", "INFO")
            return True
            
        elif os.path.isfile(img_src):
            # Local file path
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            
            try:
                from PIL import Image as PILImage
                pil_img = PILImage.open(img_src)
                img_width, img_height = pil_img.size
                width_inches = min(max_width_inches, img_width / 96.0)
                if width_inches < 3.0:
                    width_inches = min(max_width_inches, 5.0)
                run.add_picture(img_src, width=Inches(width_inches))
            except ImportError:
                run.add_picture(img_src, width=Inches(max_width_inches))
            
            sse_log(f"Embedded local image: {alt_text} ({img_src})", "INFO")
            return True
        else:
            sse_log(f"Image source not found or unsupported: {img_src[:100]}...", "WARNING")
            return False
            
    except Exception as e:
        sse_log(f"Failed to add image '{alt_text}': {e}", "WARNING")
        return False


def is_image_line(line: str) -> bool:
    """Check if a line contains a markdown image tag"""
    return bool(re.match(r'^\s*!\[([^\]]*)\]\((.+)\)\s*$', line.strip()))


def extract_image_info(line: str) -> Optional[Tuple[str, str]]:
    """Extract (alt_text, src) from a markdown image line"""
    match = re.match(r'^\s*!\[([^\]]*)\]\((.+)\)\s*$', line.strip())
    if match:
        return (match.group(1), match.group(2))
    return None


def convert_markdown_to_docx(markdown_content: str, output_path: str, title: str = None) -> str:
    """
    Convert markdown content to a Word (.docx) document with academic paper formatting.
    v41: Now properly handles base64 embedded images and local image paths.
    
    Args:
        markdown_content: The markdown text to convert
        output_path: Path to save the .docx file
        title: Optional title override
    
    Returns:
        Path to the generated .docx file
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed. Install with: pip install python-docx")
    
    doc = Document()
    setup_document_styles(doc)
    
    lines = markdown_content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    metadata_section = True  # Track if we're in the metadata header
    images_embedded = 0
    images_failed = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block - add as formatted text
                code_text = '\n'.join(code_block_lines)
                if code_text.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    # Add light gray background
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                    p._element.get_or_add_pPr().append(shading)
                code_block_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # v41: Check for image line BEFORE table check and paragraph merging
        if is_image_line(line):
            img_info = extract_image_info(line)
            if img_info:
                alt_text, img_src = img_info
                success = add_image_to_doc(doc, img_src, alt_text)
                if success:
                    images_embedded += 1
                    # Add caption below image if alt_text is meaningful
                    if alt_text and alt_text.strip():
                        cap_p = doc.add_paragraph()
                        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap_run = cap_p.add_run(alt_text)
                        cap_run.font.name = 'Times New Roman'
                        cap_run.font.size = Pt(10)
                        cap_run.italic = True
                else:
                    images_failed += 1
                    # Add placeholder text for failed images
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(f"[Image: {alt_text}]")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    run.italic = True
                    run.font.color.rgb = RGBColor(128, 128, 128)
            i += 1
            continue
        
        # Check for table
        if line.strip().startswith('|') and i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
            rows, end_idx = parse_markdown_table(lines, i)
            if rows:
                add_table_to_doc(doc, rows)
            i = end_idx
            continue
        
        # Parse the line
        line_type, segments = parse_markdown_line(line)
        
        if line_type == 'heading_1':
            # # Title - use Title style
            text = segments[0][0] if segments else ''
            if metadata_section:
                p = doc.add_paragraph(text, style='Title')
                metadata_section = False
            else:
                doc.add_heading(text, level=0)
            i += 1
            continue
        
        if line_type == 'heading_2':
            # ## Section heading
            text = segments[0][0] if segments else ''
            metadata_section = False
            doc.add_heading(text, level=1)
            i += 1
            continue
        
        if line_type == 'heading_3':
            # ### Subsection heading
            text = segments[0][0] if segments else ''
            doc.add_heading(text, level=2)
            i += 1
            continue
        
        if line_type == 'heading_4':
            # #### Sub-subsection heading
            text = segments[0][0] if segments else ''
            doc.add_heading(text, level=3)
            i += 1
            continue
        
        if line_type.startswith('heading_'):
            # Other heading levels
            text = segments[0][0] if segments else ''
            level = int(line_type.split('_')[1])
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue
        
        if line_type == 'hr':
            # Horizontal rule - skip (sections are separated by headings in Word)
            i += 1
            continue
        
        if line_type == 'list_item':
            # Add as bullet point
            add_formatted_paragraph(doc, segments, style='List Bullet')
            i += 1
            continue
        
        if line_type == 'paragraph':
            # Handle metadata lines (bold key: value pairs at the top)
            line_text = line.strip()
            if metadata_section and line_text.startswith('**') and ':' in line_text:
                # Metadata line like **Generated**: 2025-01-01
                add_formatted_paragraph(doc, segments, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                i += 1
                continue
            
            metadata_section = False
            
            # Check if this is a continuation paragraph (merge with next lines)
            # v41: Stop merging if next line is an image
            paragraph_lines = [line.strip()]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                # Stop if next line is empty, heading, list, table, hr, code block, or image
                if (not next_line or 
                    next_line.startswith('#') or 
                    next_line.startswith('|') or
                    next_line.startswith('---') or
                    next_line.startswith('***') or
                    next_line.startswith('```') or
                    is_image_line(next_line) or
                    re.match(r'^(\s*)([-*+]|\d+\.)\s+', next_line)):
                    break
                paragraph_lines.append(next_line)
                j += 1
            
            # Combine and parse the full paragraph
            full_text = ' '.join(paragraph_lines)
            segments = parse_inline_formatting(full_text)
            add_formatted_paragraph(doc, segments)
            i = j
            continue
        
        # Empty line or unrecognized
        i += 1
    
    # Log image embedding stats
    if images_embedded > 0 or images_failed > 0:
        sse_log(f"Image embedding: {images_embedded} successful, {images_failed} failed", "INFO")
    
    # Save the document
    doc.save(output_path)
    return output_path


def convert_report_to_docx(md_file_path: str, output_dir: str = None) -> Optional[str]:
    """
    Convert a markdown report file to Word (.docx) format.
    
    Args:
        md_file_path: Path to the markdown report file
        output_dir: Optional output directory (defaults to same directory as md file)
    
    Returns:
        Path to the generated .docx file, or None if conversion failed
    """
    if not DOCX_AVAILABLE:
        sse_log("python-docx not installed - Word conversion SKIPPED. Install with: pip install python-docx", "WARNING")
        sse_log("To enable Word export, run: pip install python-docx>=1.1.0", "WARNING")
        return None
    
    try:
        # Read markdown content
        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Determine output path
        if output_dir is None:
            output_dir = os.path.dirname(md_file_path)
        
        base_name = os.path.splitext(os.path.basename(md_file_path))[0]
        docx_path = os.path.join(output_dir, f"{base_name}.docx")
        
        sse_log(f"[96%] Converting report to Word document...", "INFO")
        
        # Convert
        convert_markdown_to_docx(md_content, docx_path)
        
        sse_log(f"[97%] Word document saved: {docx_path}", "SUCCESS")
        
        return docx_path
    
    except Exception as e:
        sse_log(f"Word conversion failed: {e}", "WARNING")
        return None
