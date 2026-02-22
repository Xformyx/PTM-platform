#!/usr/bin/env python3
"""Generate PTM Platform Technology Stack document as .docx"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from pathlib import Path


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(row, bg_color="1F4E79"):
    for cell in row.cells:
        set_cell_shading(cell, bg_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_header_row(table.rows[0])

    # Data
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Consolas" if ci == 1 else "Calibri"
        # Alternate row shading
        if ri % 2 == 1:
            for cell in table.rows[ri + 1].cells:
                set_cell_shading(cell, "F2F7FB")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def main():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)

    # Title
    title = doc.add_heading("PTM Analysis Platform — Technology Stack", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Version 0.1.0 | Generated for internal documentation")
    doc.add_paragraph("")

    # =====================================================================
    # A. Infrastructure / DevOps
    # =====================================================================
    doc.add_heading("A. Infrastructure / DevOps", level=1)

    add_table(doc,
        ["기술", "버전", "역할"],
        [
            ["Docker", "-", "컨테이너화 (10개 서비스)"],
            ["Docker Compose", "v2", "멀티 컨테이너 오케스트레이션"],
            ["Nginx", "Alpine", "Reverse Proxy, Rate Limiting, Static File Serving, SSL Termination"],
            ["Cloudflare Tunnel", "-", "외부 도메인(ptm.xformyx.com) → 로컬 서버 연결"],
        ],
        col_widths=[4, 2.5, 11]
    )

    # =====================================================================
    # B. Backend
    # =====================================================================
    doc.add_heading("B. Backend", level=1)

    # B-1 API Server
    doc.add_heading("B-1. API Server (api-server)", level=2)
    doc.add_paragraph("Python 3.11 (slim) 기반 FastAPI 서버. 주문 관리, 인증, 파일 업로드, 실시간 SSE, 작업 큐 발행 담당.")

    add_table(doc,
        ["카테고리", "패키지", "버전", "역할"],
        [
            ["Web Framework", "FastAPI", ">=0.115.0", "REST API 서버"],
            ["ASGI Server", "Uvicorn (standard)", ">=0.30.0", "ASGI 실행 서버 (HTTP/2, WebSocket)"],
            ["ORM", "SQLAlchemy (asyncio)", ">=2.0.0", "비동기 ORM, DB 모델"],
            ["DB Driver", "asyncmy", ">=0.2.9", "MySQL 비동기 드라이버"],
            ["Migration", "Alembic", ">=1.14.0", "DB 스키마 마이그레이션"],
            ["Validation", "Pydantic", ">=2.0.0", "요청/응답 데이터 검증"],
            ["Settings", "pydantic-settings", ">=2.0.0", "환경 변수 설정 관리"],
            ["File Upload", "python-multipart", ">=0.0.9", "Multipart form 파일 업로드"],
            ["Auth (JWT)", "python-jose (cryptography)", ">=3.3.0", "JWT 토큰 생성/검증"],
            ["Auth (Password)", "passlib (bcrypt)", ">=1.7.4", "비밀번호 해싱"],
            ["Task Queue", "Celery (redis)", ">=5.4.0", "비동기 작업 큐 발행"],
            ["Cache/PubSub", "redis", ">=5.0.0", "SSE Pub/Sub, 세션 캐시"],
            ["SSE", "sse-starlette", ">=2.0.0", "실시간 Server-Sent Events"],
            ["HTTP Client", "httpx", ">=0.27.0", "내부 서비스 간 비동기 통신"],
            ["Async File I/O", "aiofiles", ">=24.1.0", "비동기 파일 읽기/쓰기"],
            ["Data Processing", "pandas", ">=2.0.0", "config.xlsx 파싱"],
            ["Excel Reader", "openpyxl", ">=3.1.0", ".xlsx 파일 읽기"],
            ["Build System", "hatchling", "-", "pyproject.toml 기반 패키징"],
        ],
        col_widths=[3, 4.5, 2.5, 7.5]
    )

    # B-2 MCP Server
    doc.add_heading("B-2. MCP Server (mcp-server) — Bio-Database Gateway", level=2)
    doc.add_paragraph("외부 생물정보 데이터베이스 API에 대한 통합 게이트웨이. Redis 캐싱 및 Rate Limiting 적용.")

    add_table(doc,
        ["카테고리", "패키지", "버전", "역할"],
        [
            ["Web Framework", "FastAPI", ">=0.115.0", "REST API (bio-database 프록시)"],
            ["ASGI Server", "Uvicorn (standard)", ">=0.30.0", "비동기 HTTP 서버"],
            ["HTTP Client", "httpx", ">=0.27.0", "외부 API 호출 (UniProt, KEGG, STRING, PubMed, InterPro)"],
            ["Cache", "redis", ">=5.0.0", "API 응답 캐싱, Rate Limiting"],
            ["Validation", "Pydantic", ">=2.0.0", "요청/응답 모델"],
            ["Settings", "pydantic-settings", ">=2.0.0", "환경 변수 관리"],
        ],
        col_widths=[3, 4.5, 2.5, 7.5]
    )

    doc.add_paragraph("")
    p = doc.add_paragraph("연동 외부 API: ")
    p.runs[0].bold = True
    apis = [
        "NCBI PubMed (E-utilities: esearch, efetch) — 논문 검색",
        "UniProt — 단백질 정보 (기능, 위치, GO terms)",
        "KEGG — 경로(pathway) 분석",
        "STRING-DB — 단백질-단백질 상호작용",
        "InterPro — 도메인/패밀리 분류",
    ]
    for api in apis:
        doc.add_paragraph(api, style="List Bullet")

    # B-3 Workers
    doc.add_heading("B-3. Workers (Celery) — 3개 전문 워커", level=2)
    doc.add_paragraph("Preprocessing, RAG Enrichment, Report Generation 각각 독립된 Celery 워커로 실행. 멀티프로세싱 지원.")

    doc.add_heading("Core Dependencies", level=3)
    add_table(doc,
        ["카테고리", "패키지", "버전", "역할"],
        [
            ["Task Queue", "Celery (redis)", ">=5.4.0", "비동기 분산 작업 처리"],
            ["Message Broker", "redis", ">=5.0.0", "Celery 메시지 브로커"],
            ["DB (sync)", "pymysql", ">=1.1.0", "워커 → MySQL 상태 업데이트"],
            ["ORM", "SQLAlchemy", ">=2.0.0", "DB 모델 정의"],
            ["DB Driver (async)", "asyncmy", ">=0.2.9", "모델 정의 호환용"],
            ["HTTP (sync)", "requests", ">=2.31.0", "MCP/LLM 동기 호출"],
            ["HTTP (async)", "httpx", ">=0.27.0", "내부 통신"],
        ],
        col_widths=[3, 4.5, 2.5, 7.5]
    )

    doc.add_heading("Data Processing Libraries", level=3)
    add_table(doc,
        ["패키지", "버전", "역할"],
        [
            ["pandas", ">=2.0.0", "TSV/Excel 데이터 프레임 처리"],
            ["numpy", ">=1.24.0", "수치 연산, 통계"],
            ["scipy", ">=1.9.0", "통계 검정 (t-test, fold-change)"],
            ["biopython", ">=1.79", "FASTA 파싱, 서열 분석"],
            ["openpyxl", ">=3.0.0", "Excel 파일 읽기/쓰기"],
            ["lxml", ">=4.9.0", "XML 파싱 (PubMed 응답)"],
            ["tqdm", ">=4.64.0", "진행률 표시"],
        ],
        col_widths=[4, 3, 10.5]
    )

    doc.add_heading("AI/ML Libraries", level=3)
    add_table(doc,
        ["패키지", "버전", "역할"],
        [
            ["LangGraph", ">=0.2.0", "Report 생성 StateGraph 오케스트레이션 (7-node pipeline)"],
            ["langchain-core", ">=0.3.0", "LangGraph 기반 프레임워크"],
            ["ChromaDB", ">=0.5.0", "벡터 DB (RAG 문헌 검색)"],
            ["sentence-transformers", ">=3.0.0", "텍스트 임베딩 생성 (ChromaDB 색인용)"],
            ["rank-bm25", ">=0.2.2", "BM25 기반 리랭킹"],
        ],
        col_widths=[4, 3, 10.5]
    )

    doc.add_heading("Visualization / Export", level=3)
    add_table(doc,
        ["패키지", "버전", "역할"],
        [
            ["py4cytoscape", ">=1.9.0", "Cytoscape Desktop 연동 (네트워크 시각화)"],
            ["python-docx", ">=1.1.0", "Markdown → Word (.docx) 변환"],
            ["Pillow", ">=10.0.0", "이미지 처리 (docx 내 이미지 삽입)"],
        ],
        col_widths=[4, 3, 10.5]
    )

    doc.add_heading("LLM Integration", level=3)
    add_table(doc,
        ["서비스", "기본값", "역할"],
        [
            ["Ollama (로컬)", "qwen2.5:32b", "기본 LLM (리포트 생성, 가설 생성)"],
            ["OpenAI API", "(선택)", "클라우드 LLM 옵션"],
            ["Google Gemini", "(선택)", "클라우드 LLM 옵션"],
        ],
        col_widths=[4, 3, 10.5]
    )

    # B-4 Shared Infrastructure
    doc.add_heading("B-4. Shared Infrastructure (Docker Images)", level=2)
    add_table(doc,
        ["서비스", "Docker Image", "역할"],
        [
            ["MySQL", "mysql:8.0", "주 관계형 DB (주문, 사용자, 상태 관리)"],
            ["Redis", "redis:7-alpine", "Celery 브로커(DB1), Result Backend(DB2), SSE Pub/Sub(DB0), MCP 캐시(DB3)"],
            ["ChromaDB", "chromadb/chroma:latest", "벡터 DB (논문/교재 임베딩 저장, RAG 검색)"],
        ],
        col_widths=[3, 5, 9.5]
    )

    # =====================================================================
    # C. Frontend
    # =====================================================================
    doc.add_heading("C. Frontend", level=1)
    doc.add_paragraph("Node.js 20 (Alpine)에서 빌드 후 Nginx Alpine으로 정적 파일 서빙. React 18 + TypeScript + Vite 기반 SPA.")

    doc.add_heading("Production Dependencies", level=2)
    add_table(doc,
        ["카테고리", "패키지", "버전", "역할"],
        [
            ["UI Framework", "React", "^18.3.0", "SPA 컴포넌트 기반 UI"],
            ["React DOM", "react-dom", "^18.3.0", "React DOM 렌더링"],
            ["Routing", "react-router-dom", "^7.1.0", "클라이언트 사이드 라우팅"],
            ["Chart/Graph", "Plotly.js", "^2.35.0", "데이터 시각화 (차트, 그래프)"],
            ["Chart (React)", "react-plotly.js", "^2.6.0", "React용 Plotly 래퍼"],
            ["State Management", "Zustand", "^5.0.0", "경량 전역 상태 관리"],
            ["CSS Utility", "clsx", "^2.1.0", "조건부 CSS 클래스 결합"],
            ["CSS Utility", "tailwind-merge", "^2.6.0", "Tailwind 클래스 충돌 해결"],
        ],
        col_widths=[3, 4.5, 2.5, 7.5]
    )

    doc.add_heading("Dev Dependencies", level=2)
    add_table(doc,
        ["패키지", "버전", "역할"],
        [
            ["TypeScript", "^5.6.0", "정적 타입 체크 (ES2020 target)"],
            ["Vite", "^6.0.0", "빌드 도구 (HMR, 번들링)"],
            ["@vitejs/plugin-react", "^4.3.0", "Vite React 플러그인"],
            ["Tailwind CSS", "^3.4.0", "유틸리티 퍼스트 CSS 프레임워크"],
            ["PostCSS", "^8.4.0", "CSS 후처리 파이프라인"],
            ["Autoprefixer", "^10.4.0", "CSS 벤더 프리픽스 자동 추가"],
            ["@types/react", "^18.3.0", "React 타입 정의"],
            ["@types/react-dom", "^18.3.0", "React DOM 타입 정의"],
        ],
        col_widths=[4, 3, 10.5]
    )

    # =====================================================================
    # D. Architecture Diagram
    # =====================================================================
    doc.add_heading("D. 서비스 아키텍처", level=1)

    diagram = """Client (Browser)
    │
    ▼
┌─────────────┐
│   Nginx     │ :80/:443  (Gateway)
│  (Alpine)   │
└──────┬──────┘
       │
  ┌────┴────┐
  │         │
  ▼         ▼
Frontend   API Server
(React)    (FastAPI)
 :80        :8000
            │
     ┌──────┼──────────┐
     │      │          │
     ▼      ▼          ▼
   MySQL   Redis     MCP Server
   :3306   :6379     (FastAPI) :8001
            │          │
            │          └─→ UniProt, KEGG, STRING, PubMed, InterPro
            │
     ┌──────┼──────────┐
     │      │          │
     ▼      ▼          ▼
  Worker   Worker    Worker
 (Preproc) (RAG)   (Report)
     │       │         │
     │       │         ├─→ ChromaDB :8000
     │       │         ├─→ Ollama (host) :11434
     └───────┴─────────┴─→ Cytoscape (host) :1234"""

    p = doc.add_paragraph()
    run = p.add_run(diagram)
    run.font.name = "Consolas"
    run.font.size = Pt(8)

    # =====================================================================
    # E. Summary
    # =====================================================================
    doc.add_heading("E. 총 의존성 수 요약", level=1)

    add_table(doc,
        ["영역", "Python 패키지", "JS 패키지", "Docker 이미지"],
        [
            ["API Server", "14", "-", "1 (python:3.11-slim)"],
            ["MCP Server", "6", "-", "1 (python:3.11-slim)"],
            ["Workers", "22", "-", "1 (python:3.11-slim)"],
            ["Frontend", "-", "7 prod + 8 dev", "2 (node:20-alpine → nginx:alpine)"],
            ["Infrastructure", "-", "-", "3 (mysql:8.0, redis:7-alpine, chromadb/chroma)"],
            ["합계", "42", "15", "8"],
        ],
        col_widths=[4, 3.5, 4.5, 5.5]
    )

    # Save
    output_path = Path(__file__).parent / "PTM_Platform_Tech_Stack.docx"
    doc.save(str(output_path))
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    main()
